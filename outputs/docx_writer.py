"""
DOCX output utilities.

Creates bilingual curriculum documents.
"""


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH



def save_docx(
    resource,
    output_path,
    title="English - isiZulu Translation"
):
    """
    Generate a bilingual Word document.

    Parameters:

        resource (dict):
            Translation resource

        output_path (str):
            Destination DOCX file

        title (str):
            Document title
    """
    #translations = resource["content"]

    document = Document()


    # Title

    heading = document.add_heading(
        title,
        level=1
    )

    heading.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )


    document.add_paragraph(
        "English and isiZulu curriculum translation"
    )


    document.add_page_break()



    # Content

    for section in resource["sections"]:

        document.add_heading(
            f"Section {section['section_id']}",
            level=2
        )


        document.add_paragraph(
            "🇿🇦 isiZulu Learning Content",
            style="Intense Quote"
        )


        document.add_paragraph(
            section["lesson_content"]
        )


        document.add_paragraph(
            "🇬🇧 English Reference",
            style="Intense Quote"
        )


        document.add_paragraph(
            section["english_reference"]
        )


        if section["key_terms"]:

            document.add_heading(
                "📘 Key Terms",
                level=3
            )


            for term in section["key_terms"]:

                document.add_paragraph(
                    f"{term['zulu_term']} "
                    f"({term['term']})"
                )


                document.add_paragraph(
                    term["meaning"]
                )


        document.add_paragraph(
            "--------------------------------"
        )


    document.save(
        output_path
    )
    
    return output_path